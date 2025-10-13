"""
    TemplateGofer
    --------------
    A desktop utility that automates find-and-replace operations in Word documents
    using mappings defined in a spreadsheet file.

    The user is prompted (via file dialogs) to select:
      1. A spreadsheet file (.xlsx or .csv) containing two columns:
         - 'find'    → the text to search for
         - 'replace' → the text to replace it with
      2. A Word document (.docx) file to apply replacements on.
"""
# custom modules
from kbtoolbox import dir_filedialog
from tkinter import filedialog, messagebox
import os
import pandas as pd
from docx import Document
import sys
import subprocess
from openpyxl import Workbook


class TemplateGofer:
    """
        TemplateGofer:
      • Reads the spreadsheet and extracts all find/replace pairs.
      • Opens the selected Word document.
      • Replaces every instance of each "find" string with its corresponding "replace" string
        in both paragraphs and table cells.
      • Saves a new copy of the modified document in the same directory, appending "_tg"
        to the filename (e.g., "example.docx" → "example_tg.docx").
      • Displays a success message if the process completes, or an error message otherwise.
    """
    def __init__(self):
        self.frame = None
        self.spreadsheet_path = None
        self.doc_path = None
        self.replacements = []
        self.new_file = None

    def select_spreadsheet(self):
        """Prompt user to select the spreadsheet file."""
        self.spreadsheet_path = filedialog.askopenfilename(
            title="Select Spreadsheet File",
            initialdir=dir_filedialog(),
            filetypes=[("Spreadsheet Files", "*.xlsx *.csv")]
        )
        if not self.spreadsheet_path:
            raise FileNotFoundError("No spreadsheet file selected.")

    def load_replacements(self):
        """Load find/replace pairs from the spreadsheet."""
        if self.spreadsheet_path.lower().endswith('.csv'):
            df = pd.read_csv(self.spreadsheet_path)
        else:
            df = pd.read_excel(self.spreadsheet_path)

        df.columns = [c.lower().strip() for c in df.columns]
        if not {'find', 'replace'}.issubset(df.columns):
            raise ValueError("Spreadsheet must have 'find' and 'replace' columns.")

        self.replacements = list(zip(df['find'].astype(str), df['replace'].astype(str)))

    def select_doc_file(self):
        """Prompt user to select the Word file."""
        self.doc_path = filedialog.askopenfilename(
            title="Select Word Document File",
            initialdir=dir_filedialog(),
            filetypes=[("Word Files", "*.docx")]
        )
        if not self.doc_path:
            raise FileNotFoundError("No Word document selected.")

    @staticmethod
    def check_file_in_use(filepath):
        """
        Check if the given file is currently locked or in use by another process.
        Returns True if the file is locked.
        """
        if not os.path.exists(filepath):
            return False
        # --- Check LibreOffice lock file ---
        folder = os.path.dirname(filepath)
        filename = os.path.basename(filepath)
        libre_lock = os.path.join(folder, f".~lock.{filename}#")
        if os.path.exists(libre_lock):
            return True
        # check if open in Microsoft Word
        try:
            with open(filepath, 'a'):
                pass
            return False
        except PermissionError:  # if it is already open, there will be a permission error
            return True

    def process_document(self):
        """Perform find/replace and save modified file."""
        base, ext = os.path.splitext(self.doc_path)  # name new file
        new_path = f"{base}_tg{ext}"
        if self.check_file_in_use(new_path):
            raise PermissionError(
                f"The file '{os.path.basename(self.doc_path)}' is currently in use.\n"
                f"Please close it and try again."
            )
        doc = Document(self.doc_path)
        for find_text, replace_text in self.replacements:
            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                if find_text in paragraph.text:
                    for run in paragraph.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find_text in cell.text:
                            cell.text = cell.text.replace(find_text, replace_text)
        # Save as new file
        try:
            doc.save(new_path)
        except PermissionError:
            raise PermissionError(
                f"Unable to save '{os.path.basename(new_path)}'.\n"
                f"It may be open in another program."
            )

        return new_path

    def open(self):
        """ open the spreadsheet """
        try:
            if sys.platform == "win32":  # open the text document
                os.startfile(self.new_file)
            if sys.platform == "linux":
                subprocess.call(["xdg-open", 'kb_sub/infc_grv/' + os.path.basename(self.new_file)])
            if sys.platform == "darwin":
                subprocess.call(["open", self.new_file])
        except PermissionError:
            messagebox.showerror("Template Gofer",
                                 "The process has failed and the document will not open.",
                                 parent=self.frame)

    def run(self, frame):
        """Run the Template Gofer process with dialogs and error handling."""
        self.frame = frame
        try:
            self.select_spreadsheet()
            self.load_replacements()
            self.select_doc_file()
            self.new_file = self.process_document()
            messagebox.showinfo("Template Gofer",
                                f"Template Gofer completed successfully!\n\nSaved as:\n{self.new_file}",
                                parent=self.frame)
            self.open()

        except Exception as e:
            messagebox.showerror("Template Gofer", f"Process failed and is canceled.\n\n{str(e)}",
                                 parent=self.frame)


class GenTgSpreadsheet:
    """
    A class to create a basic XLSX file template with 'find' and 'replace'
    headers, allowing the user to select the save location using a file dialog.
    """

    def __init__(self):
        self.frame = None
        self.filepath = None
        self.workbook = None
        self.sheet = None

    def run(self, frame):
        """ Prompts the user to select a save location and then creates and saves an XLSX file with pre-defined
        headers in columns A and B. """
        self.frame = frame
        if not self._name_file():
            return
        self._generate_workbook()
        self._write_headers()
        self._saveandopen()

    def _name_file(self):
        """ name and choose the location of the file """
        self.filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            initialdir=dir_filedialog(),
            filetypes=[("Excel files", "*.xlsx")],
            title="Choose location and name for the Template Gofer Spreadsheet"
        )
        if not self.filepath:  # The user cancelled the dialog
            return False
        return True

    def _generate_workbook(self):
        """ generate the workbook """
        self.workbook = Workbook()
        self.sheet = self.workbook.active
        self.sheet.title = "Template Gofer"

    def _write_headers(self):
        # Define and write the headers in the first row (A1 and B1)
        self.sheet['A1'] = "find"
        self.sheet['B1'] = "replace"
        # Adjust column widths for better visibility
        self.sheet.column_dimensions['A'].width = 30
        self.sheet.column_dimensions['B'].width = 30

    def _saveandopen(self):
        """ save and open the spreadsheet """
        self.workbook.save(self.filepath)  # Save the workbook to the selected file path
        try:
            self.workbook.save(self.filepath)
            if sys.platform == "win32":
                os.startfile(self.filepath)
            if sys.platform == "linux":
                subprocess.call(["xdg-open", 'kb_sub/infc_grv/' + os.path.basename(self.filepath)])
            if sys.platform == "darwin":
                subprocess.call(["open", self.filepath])
        except PermissionError:
            messagebox.showerror("Template Gofer",
                                 "The process has failed and the spreadsheet will not open.",
                                 parent=self.frame)
